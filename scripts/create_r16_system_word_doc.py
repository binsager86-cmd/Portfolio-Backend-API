from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "r16_system_explanation"
ASSET_DIR = OUT_DIR / "assets"
DOCX_PATH = OUT_DIR / "R16_Eagle_Eye_V5_System_Explanation.docx"

W = 2400
BG = "#f7f8fb"
INK = "#152238"
MUTED = "#5d697b"
BLUE = "#dcecff"
GREEN = "#dff5e8"
YELLOW = "#fff4ce"
RED = "#ffe0df"
LAV = "#ece6ff"
LINE = "#526173"


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        r"C:\Windows\Fonts\segoeuib.ttf" if bold else r"C:\Windows\Fonts\segoeui.ttf",
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
    ]
    for path in candidates:
        if Path(path).exists():
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


TITLE = font(54, True)
SUB = font(34, True)
BODY = font(28)
SMALL = font(23)
DIAGRAM_SMALL = font(21)


def wrap_text(draw: ImageDraw.ImageDraw, text: str, fnt: ImageFont.ImageFont, max_width: int) -> list[str]:
    lines: list[str] = []
    for raw in text.split("\n"):
        words = raw.split()
        if not words:
            lines.append("")
            continue
        line = words[0]
        for word in words[1:]:
            candidate = f"{line} {word}"
            if draw.textbbox((0, 0), candidate, font=fnt)[2] <= max_width:
                line = candidate
            else:
                lines.append(line)
                line = word
        lines.append(line)
    return lines


def box(draw: ImageDraw.ImageDraw, xy: tuple[int, int, int, int], title: str, body: str = "", fill: str = BLUE) -> None:
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=28, fill=fill, outline="#9aa8ba", width=4)
    draw.text((x1 + 28, y1 + 22), title, fill=INK, font=SUB)
    if body:
        y = y1 + 76
        for line in wrap_text(draw, body, BODY, x2 - x1 - 56):
            draw.text((x1 + 28, y), line, fill=MUTED, font=BODY)
            y += 36


def small_box(draw: ImageDraw.ImageDraw, xy: tuple[int, int, int, int], title: str, body: str = "", fill: str = BLUE) -> None:
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=24, fill=fill, outline="#9aa8ba", width=4)
    draw.text((x1 + 24, y1 + 18), title, fill=INK, font=BODY)
    if body:
        y = y1 + 60
        for line in wrap_text(draw, body, DIAGRAM_SMALL, x2 - x1 - 48):
            draw.text((x1 + 24, y), line, fill=MUTED, font=DIAGRAM_SMALL)
            y += 28


def label(draw: ImageDraw.ImageDraw, xy: tuple[int, int], text: str, fnt: ImageFont.ImageFont = BODY, color: str = INK) -> None:
    draw.text(xy, text, fill=color, font=fnt)


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str = LINE, width: int = 8) -> None:
    draw.line([start, end], fill=color, width=width)
    sx, sy = start
    ex, ey = end
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex >= sx else -1
        pts = [(ex, ey), (ex - direction * 34, ey - 20), (ex - direction * 34, ey + 20)]
    else:
        direction = 1 if ey >= sy else -1
        pts = [(ex, ey), (ex - 20, ey - direction * 34), (ex + 20, ey - direction * 34)]
    draw.polygon(pts, fill=color)


def poly_arrow(draw: ImageDraw.ImageDraw, points: list[tuple[int, int]], color: str = LINE, width: int = 8) -> None:
    for start, end in zip(points, points[1:]):
        draw.line([start, end], fill=color, width=width)
    start = points[-2]
    end = points[-1]
    sx, sy = start
    ex, ey = end
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex >= sx else -1
        pts = [(ex, ey), (ex - direction * 34, ey - 20), (ex - direction * 34, ey + 20)]
    else:
        direction = 1 if ey >= sy else -1
        pts = [(ex, ey), (ex - 20, ey - direction * 34), (ex + 20, ey - direction * 34)]
    draw.polygon(pts, fill=color)


def canvas(height: int, title: str) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    img = Image.new("RGB", (W, height), BG)
    draw = ImageDraw.Draw(img)
    draw.text((80, 54), title, fill=INK, font=TITLE)
    draw.line([(80, 126), (W - 80, 126)], fill="#c8d0da", width=4)
    return img, draw


def save(img: Image.Image, name: str) -> Path:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    path = ASSET_DIR / name
    img.save(path, "PNG", optimize=True, dpi=(220, 220))
    return path


def diagram_architecture() -> Path:
    img, d = canvas(1500, "R16 v5 System Architecture")
    box(d, (90, 230, 470, 420), "Sealed DB", "OHLCV\nindicators\nsegments", YELLOW)
    box(d, (620, 210, 1060, 450), "Layer 1 Harness", "Reads data\ncalls engines\nwrites evidence", BLUE)
    box(d, (1220, 170, 1740, 330), "Warmup + Adapter", "Readiness and segment-safe daily payloads", GREEN)
    box(d, (1220, 390, 1740, 550), "Base Geometry", "Base forming, valid, retired, invalidated", GREEN)
    box(d, (1220, 610, 1740, 770), "Flow Confirmation", "Intent and confirmation state", GREEN)
    box(d, (1220, 830, 1740, 990), "Pivot Engine", "3-session lag and significant pivots", GREEN)
    box(d, (1840, 455, 2300, 705), "Session Context", "One complete daily fact packet sent to pure logic", LAV)
    box(d, (620, 1080, 1120, 1325), "Layer 2 State Machine", "Pure transition: state + context -> new state + actions", RED)
    box(d, (1360, 1090, 1780, 1315), "Actions", "OPEN_POSITION\nCLOSE_POSITION\nDAILY_STATE", YELLOW)
    box(d, (1910, 1080, 2300, 1325), "Evidence", "Harness DB\nexports\nSHA sidecars", BLUE)
    arrow(d, (470, 325), (620, 325))
    arrow(d, (1060, 325), (1220, 250))
    arrow(d, (1060, 340), (1220, 470))
    arrow(d, (1060, 355), (1220, 690))
    arrow(d, (1060, 370), (1220, 910))
    arrow(d, (1740, 250), (1840, 545))
    arrow(d, (1740, 470), (1840, 570))
    arrow(d, (1740, 690), (1840, 595))
    arrow(d, (1740, 910), (1840, 620))
    poly_arrow(d, [(2070, 705), (2070, 1030), (870, 1030), (870, 1080)])
    arrow(d, (1120, 1205), (1360, 1205))
    arrow(d, (1780, 1205), (1910, 1205))
    label(d, (80, 1410), "Key idea: existing ratified engines describe the market; R16 only decides lifecycle and trade actions.", BODY, MUTED)
    return save(img, "01_architecture.png")


def diagram_daily_flow() -> Path:
    img, d = canvas(980, "One Trading Day Replay")
    items = [
        ("1. Load Day", "price, volume, indicators"),
        ("2. Normalize", "calendar, segment, mask"),
        ("3. Readiness", "is history sufficient?"),
        ("4. Base", "forming, valid, retired"),
        ("5. Flow", "intent + confirmation"),
        ("6. Context", "single fact packet"),
        ("7. Step", "state machine action"),
        ("8. Record", "DB row + events"),
    ]
    x = 80
    y = 260
    for i, (t, b) in enumerate(items):
        bx = x + (i % 4) * 560
        by = y + (i // 4) * 330
        box(d, (bx, by, bx + 390, by + 170), t, b, [YELLOW, GREEN, GREEN, GREEN, GREEN, LAV, RED, BLUE][i])
        if i % 4 != 3:
            arrow(d, (bx + 390, by + 85), (bx + 560, by + 85))
        elif i == 3:
            arrow(d, (bx + 195, by + 170), (bx + 195, by + 330))
    arrow(d, (1850, 675), (640, 675))
    return save(img, "02_daily_flow.png")


def diagram_lifecycle() -> Path:
    img, d = canvas(1320, "Lifecycle State Machine")
    nodes = {
        "NEUTRAL": (120, 300, 460, 440, BLUE),
        "BASE_FORMING": (610, 300, 1010, 440, BLUE),
        "BASE_VALID": (1160, 300, 1540, 440, GREEN),
        "MARKUP_ACTIVE": (1710, 300, 2190, 440, GREEN),
        "AVOID_SOFT": (1160, 650, 1540, 800, YELLOW),
        "AVOID_HARD": (1710, 650, 2190, 800, RED),
        "MARKDOWN": (1710, 980, 2190, 1130, RED),
    }
    for name, (x1, y1, x2, y2, color) in nodes.items():
        box(d, (x1, y1, x2, y2), name, "", color)
    arrow(d, (460, 370), (610, 370))
    arrow(d, (1010, 370), (1160, 370))
    arrow(d, (1540, 370), (1710, 370))
    arrow(d, (1350, 440), (1350, 650))
    arrow(d, (1540, 725), (1710, 725))
    arrow(d, (1950, 800), (1950, 980))
    arrow(d, (1710, 1055), (460, 1055))
    label(d, (550, 460), "base begins", SMALL, MUTED)
    label(d, (1050, 460), "valid base", SMALL, MUTED)
    label(d, (1545, 460), "upward retirement", SMALL, MUTED)
    label(d, (930, 610), "2+ soft deterioration signals", SMALL, MUTED)
    label(d, (1570, 610), "hard risk breach", SMALL, MUTED)
    label(d, (1790, 920), "risk-off / below EMA30", SMALL, MUTED)
    label(d, (700, 1000), "recovery resets to neutral", SMALL, MUTED)
    return save(img, "03_lifecycle.png")


def diagram_entry_exit() -> Path:
    img, d = canvas(1500, "Entry and Exit Decisions")
    box(d, (100, 230, 510, 420), "No Position", "candidate is flat", BLUE)
    box(d, (710, 170, 1200, 360), "Base Entry", "valid base + intent formed + flow confirmed", GREEN)
    box(d, (710, 460, 1200, 650), "Markup Entry", "pullback to EMA band or flag breakout", GREEN)
    box(d, (1430, 310, 1880, 510), "Open Position", "entry close, max close, MFE, sessions held", YELLOW)
    box(d, (100, 900, 510, 1110), "Shared Exits", "AVOID_HARD or stagnant 60-session time stop", RED)
    box(d, (710, 830, 1200, 1020), "Variant A", "exit after 2 closes below EMA30", LAV)
    box(d, (710, 1110, 1200, 1300), "Variant B", "exit below max close - 2.75 x ATR14", LAV)
    box(d, (1430, 980, 1880, 1190), "Close or Hold", "record exit reason or continue", BLUE)
    arrow(d, (510, 325), (710, 265))
    arrow(d, (510, 325), (710, 555))
    arrow(d, (1200, 265), (1430, 410))
    arrow(d, (1200, 555), (1430, 410))
    arrow(d, (1655, 510), (305, 900))
    arrow(d, (510, 1005), (710, 925))
    arrow(d, (510, 1005), (710, 1205))
    arrow(d, (1200, 925), (1430, 1080))
    arrow(d, (1200, 1205), (1430, 1080))
    label(d, (1040, 1370), "A and B share entries, avoid logic, and time-stop. Only the open-position exit style differs.", BODY, MUTED)
    return save(img, "04_entry_exit.png")


def diagram_variants() -> Path:
    img, d = canvas(1080, "Variant A vs Variant B")
    box(d, (130, 240, 1040, 500), "Variant A: Structural EMA Exit", "Question: did the medium trend break?\nRule: close below EMA30 for 2 consecutive sessions.\nExit reason: EXIT_STRUCTURAL_EMA30_2C", GREEN)
    box(d, (1360, 240, 2270, 500), "Variant B: Chandelier Exit", "Question: did price fall too far from its best close?\nRule: close < max close since entry - 2.75 x ATR14.\nExit reason: EXIT_CHANDELIER", YELLOW)
    box(d, (580, 720, 1820, 930), "Shared Foundation", "same sealed data, same ratified engines, same entries, same AVOID_HARD, same progress time-stop", BLUE)
    arrow(d, (1040, 370), (1360, 370))
    arrow(d, (800, 500), (950, 720))
    arrow(d, (1650, 500), (1450, 720))
    label(d, (1045, 330), "only exit policy differs", SMALL, MUTED)
    return save(img, "05_variants.png")


def diagram_exports() -> Path:
    img, d = canvas(1180, "Evidence and Export Package")
    box(d, (100, 240, 560, 430), "Variant A DB", "r16 daily rows\nposition events\nmetadata", BLUE)
    box(d, (100, 610, 560, 800), "Variant B DB", "same tables\nseparate replay", BLUE)
    box(d, (820, 250, 1320, 430), "Daily Evidence", "SANAM\nTIJARA\nMABANEE", GREEN)
    box(d, (820, 560, 1320, 740), "Universe Files", "per-symbol totals\nglobal aggregate", GREEN)
    box(d, (1530, 230, 2210, 410), "v5_vs_baseline", "candidate vs v4.1-B baseline", YELLOW)
    box(d, (1530, 500, 2210, 680), "Audit + Gates", "D4 evidence\nG1-G7 PASS/FAIL", YELLOW)
    box(d, (1530, 770, 2210, 950), "SHA Sidecars", "hash every exported file", RED)
    arrow(d, (560, 335), (820, 335))
    arrow(d, (560, 705), (820, 650))
    arrow(d, (1320, 335), (1530, 320))
    arrow(d, (1320, 650), (1530, 590))
    arrow(d, (1870, 680), (1870, 770))
    return save(img, "06_exports.png")


def diagram_indicators() -> Path:
    img, d = canvas(1520, "Indices and Inputs Used by R16")
    box(d, (90, 220, 560, 460), "Price Inputs", "open, high, low, close\nvolume, value_kwd", YELLOW)
    box(d, (90, 540, 560, 835), "Trend Indices", "EMA10\nEMA30\nEMA30 slope\nSMA200 context\nfrom earlier module evidence", GREEN)
    box(d, (90, 930, 560, 1210), "Volatility", "ATR14\nbase width\nATR-scaled pivots\nchandelier distance", BLUE)
    box(d, (760, 220, 1260, 500), "Flow / Participation", "OBV\nOBV slope 40\nANV slope 40\nCMF10\nrelative volume", LAV)
    box(d, (760, 570, 1260, 830), "Regime / Strength", "RSI14\nADX19\nliquidity gates\nreadiness state", GREEN)
    box(d, (760, 930, 1260, 1210), "Structure", "base top / low\nbase validity\nbase MFE\nsignificant pivots", BLUE)
    box(d, (1530, 500, 2250, 900), "SessionContext", "date, close, high, low\nbase_state, base_valid\nconfirmation_state, intent_state\nema10, ema30, atr14, obv\nusable pivots, flag_breakout", RED)
    for y in (340, 690, 1070):
        arrow(d, (560, y), (760, y))
    arrow(d, (1260, 340), (1530, 600))
    arrow(d, (1260, 700), (1530, 700))
    arrow(d, (1260, 1070), (1530, 800))
    label(d, (120, 1340), "Only Layer 1 reads these raw values. Layer 2 sees the prepared SessionContext and stays pure.", BODY, MUTED)
    return save(img, "07_indicators.png")


def diagram_trigger_matrix() -> Path:
    img, d = canvas(1740, "Lifecycle Trigger Matrix")
    rows = [
        ("BASE_FORMING", "base_state = BASE_FORMING\nwhile prior lifecycle is NEUTRAL/BASE_FORMING", BLUE),
        ("BASE_VALID", "base_reference validity = VALID\nfrom NEUTRAL / BASE_FORMING / RE_BASE", GREEN),
        ("MARKUP_ACTIVE", "base_state = BASE_RETIRED\nand base_mfe >= 20%", GREEN),
        ("AVOID_SOFT", "2 of 3 soft signals are true\nS1 lower high, S2 EMA30 loss, S3 OBV divergence", YELLOW),
        ("AVOID_HARD", "H1 close < markup swing low\nor H2 two closes below base top", RED),
        ("MARKDOWN", "after AVOID_HARD when close < EMA30", RED),
        ("NEUTRAL RESET", "MARKDOWN recovers after 5 closes above EMA30\nor retired base lacks upward MFE", BLUE),
    ]
    y = 210
    for state, rule, color in rows:
        small_box(d, (110, y, 610, y + 165), state, "", color)
        small_box(d, (760, y, 2260, y + 165), "Trigger", rule, color)
        arrow(d, (610, y + 82), (760, y + 82))
        y += 205
    return save(img, "08_trigger_matrix.png")


def diagram_calculations() -> Path:
    img, d = canvas(1340, "Key Calculations")
    box(d, (90, 220, 760, 430), "Base MFE", "base_mfe = max(0, high / base_top_ref - 1)\nMARKUP_ACTIVE threshold = 0.20", GREEN)
    box(d, (90, 520, 760, 760), "Pivot Significance", "pivot high/low forms with k=3 sessions on both sides\nusable after another 3 sessions\nsignificant if price distance >= 1.5 x ATR14", BLUE)
    box(d, (90, 850, 760, 1080), "Position MFE", "mfe = max(previous_mfe, close / entry_close - 1)\nused by progress time-stop waiver", YELLOW)
    box(d, (1050, 220, 2260, 430), "Variant A Exit", "ema30_loss_streak += 1 when close < EMA30\nexit when streak >= 2", LAV)
    box(d, (1050, 520, 2260, 760), "Variant B Exit", "trail = max_close_since_entry - 2.75 x ATR14\nexit when close < trail", LAV)
    box(d, (1050, 850, 2260, 1080), "Time Stop", "first check at 60 sessions, then every 20 sessions\nexit if MFE < 8% and lifecycle is not MARKUP_ACTIVE", RED)
    arrow(d, (760, 325), (1050, 325))
    arrow(d, (760, 970), (1050, 970))
    return save(img, "09_calculations.png")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(21, 34, 56)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)
    for run in p.runs:
        run.font.size = Pt(10.5)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item).font.size = Pt(10.5)


def add_rule_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        run = hdr[i].paragraphs[0].add_run(header)
        run.bold = True
        run.font.size = Pt(9.5)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            run.font.size = Pt(8.5)
    doc.add_paragraph()


def add_formula(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(21, 34, 56)


def add_diagram(doc: Document, path: Path, caption: str) -> None:
    doc.add_picture(str(path), width=Inches(6.7))
    p = doc.add_paragraph(caption)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(93, 105, 123)


def build_doc(diagrams: list[Path]) -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("R16 Eagle Eye v5 Candidate System")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(21, 34, 56)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Detailed rulebook: indices, lifecycle triggers, formulas, Variant A/B exits, and evidence flow")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(93, 105, 123)

    add_heading(doc, "Purpose and Scope", 1)
    add_body(doc, "This document explains the R16 Eagle Eye v5 candidate system as a full technical rulebook, not only as an executive summary. It describes the data inputs, indices, lifecycle states, triggers, calculations, entries, exits, Variant A, Variant B, and the evidence files produced by the harness.")
    add_body(doc, "R16 v5 is a controlled A/B replay over the same sealed market surface used by the ratified baseline. Both variants use the same data, same ratified services, same entry logic, same avoid logic, and same progress time-stop. The only intended A/B difference is the open-position exit rule: Variant A exits on EMA30 structural failure, while Variant B exits on a volatility-adjusted chandelier stop.")

    add_heading(doc, "System Architecture", 1)
    add_body(doc, "The architecture separates side effects from candidate logic. Layer 1 is the harness: it reads the sealed database, calls ratified services, builds session context, writes the harness database, and exports evidence. Layer 2 is pure candidate logic: it receives context and returns updated state plus actions.")
    add_diagram(doc, diagrams[0], "Figure 1. R16 v5 architecture and data flow.")

    add_heading(doc, "What Each Layer Is Responsible For", 2)
    add_rule_table(
        doc,
        ["Layer", "Allowed Responsibilities", "Not Allowed / Governance"],
        [
            ["Layer 1 Harness", "Read sealed DB; load symbol windows; call WarmupReadinessEngine, DataSurfaceAdapter, AdaptiveBaseGeometry, FlowConfirmationEngine; compute pivots and flag breakout; write DB and exports.", "Must not silently change sealed input or ratified engine behavior."],
            ["Layer 2 State Machine", "Receive one daily SessionContext; update lifecycle state; open/close/hold; emit DAILY_STATE actions.", "No DB, no files, no service imports, no clock, no symbol-specific strings."],
        ],
    )

    add_heading(doc, "Daily Replay", 1)
    add_body(doc, "For each symbol and each trading day, the harness normalizes the daily data, checks readiness, evaluates base geometry and flow confirmation, computes usable pivots, and packages all facts into a SessionContext. The state machine does not fetch anything itself; it only reacts to that context.")
    add_diagram(doc, diagrams[1], "Figure 2. One trading day replayed through the system.")

    add_heading(doc, "Indices and Inputs Used", 1)
    add_body(doc, "The candidate layer uses a prepared daily context. The raw price and indicator values are read by the harness and ratified engines. The candidate state machine sees only the cleaned fields needed for lifecycle and trade decisions.")
    add_diagram(doc, diagrams[6], "Figure 7. Indices and inputs used to build SessionContext.")
    add_rule_table(
        doc,
        ["Input / Index", "Where It Comes From", "How R16 Uses It"],
        [
            ["close", "Daily OHLCV row", "Main decision price for EMA checks, base-top breaches, swing-low breaches, chandelier stop, PnL, and MFE."],
            ["high", "Daily OHLCV row", "Used in base_mfe calculation: high / base_top_ref - 1."],
            ["low", "Daily OHLCV row", "Used for pullback entry: low touches EMA band or falls to/below EMA30."],
            ["EMA10", "Indicator payload", "Pullback recovery trigger: close must recover to at least EMA10 within 3 sessions after touch."],
            ["EMA30", "Indicator payload", "Trend boundary for Variant A, AVOID_SOFT S2, MARKDOWN transition, and recovery."],
            ["EMA30 slope", "Indicator payload as ema30_slope", "Used by soft condition S2: close below EMA30 and EMA30 slope < 0."],
            ["ATR14", "Indicator payload atr_14, fallback high-low", "Used for significant pivot filtering and Variant B chandelier distance."],
            ["OBV", "Indicator payload", "Stored at significant high pivots to detect flow divergence."],
            ["base_state", "AdaptiveBaseGeometry", "Controls NEUTRAL, BASE_FORMING, BASE_VALID, and MARKUP_ACTIVE transitions."],
            ["base_reference", "AdaptiveBaseGeometry", "Provides base validity, base top, base low, and base_reference_id."],
            ["confirmation_state", "FlowConfirmationEngine", "Entry requires CONFIRMED for direct base entry."],
            ["candidate_intent_state", "FlowConfirmationEngine candidate_intent", "Entry requires INTENT_FORMED for direct base entry."],
            ["usable pivots", "Harness PivotEngine", "Used for AVOID_SOFT S1/S3 and AVOID_HARD swing-low breach."],
            ["flag_breakout", "Harness flag detector", "Markup entry path when price breaks out of a compact 5-15 session range."],
        ],
    )

    add_heading(doc, "Ratified Engine Parameters Used by Harness", 1)
    add_body(doc, "The following parameters are passed to ratified engines by the harness. They are not tuned inside the pure candidate state machine.")
    add_rule_table(
        doc,
        ["Engine", "Parameters / Constants", "Purpose"],
        [
            ["WarmupReadinessEngine", "long lookback 180 sessions; segment restart 20 sessions; fallback 60 sessions", "Determines whether the day has enough history to evaluate signals."],
            ["AdaptiveBaseGeometry", "BASE_MIN_SESSIONS=10; BASE_MAX_WIDTH_PCT=0.24; ATR_SQUEEZE_PCTILE=0.95", "Detects base state, validity, retirement, and invalidation."],
            ["FlowConfirmationEngine", "OBV slope min 0.10; ANV slope min 0.10; CMF floor 0.05; relative volume context min 2.5; RSI regime 50; ADX trigger 18; chase band 0.08; liquidity gates 100,000 / 50,000 KWD", "Detects candidate intent and confirmation state."],
            ["PivotEngine", "confirmation lag k=3; significant pivot threshold 1.5 x ATR14", "Creates lag-safe significant highs/lows and markup swing lows."],
        ],
    )

    add_heading(doc, "Lifecycle States", 1)
    add_body(doc, "The lifecycle state is the system's compact description of where the stock is in its technical journey. It can move from neutral, to base forming, to base valid, to markup active. If evidence deteriorates, it can enter avoid soft or avoid hard. Avoid hard is the strict risk-off state and forces an open position to close.")
    add_diagram(doc, diagrams[2], "Figure 3. Conceptual lifecycle state machine.")

    add_diagram(doc, diagrams[7], "Figure 8. Exact lifecycle trigger matrix implemented in the candidate state machine.")

    add_heading(doc, "How Lifecycle Is Identified", 2)
    add_rule_table(
        doc,
        ["Lifecycle", "Trigger", "Based On", "Effect"],
        [
            ["NEUTRAL", "Initial state; or base retired without sufficient upward MFE; or MARKDOWN recovery completes.", "state memory, base_state, EMA30 recovery", "No position unless an entry condition later appears."],
            ["BASE_FORMING", "base_state is BASE_FORMING while previous lifecycle is NEUTRAL or BASE_FORMING.", "AdaptiveBaseGeometry output", "System tracks structure but does not enter yet."],
            ["BASE_VALID", "base_reference validity is VALID while lifecycle is NEUTRAL, BASE_FORMING, or RE_BASE.", "base_reference.base_validity_state", "Captures base_top_ref, base_low_ref, base_reference_id; direct entry becomes possible."],
            ["MARKUP_ACTIVE", "base_state is BASE_RETIRED and base_mfe >= 20%.", "base_state plus base_mfe formula", "Allows markup pullback and flag breakout entries; helps waive stagnant time-stop."],
            ["AVOID_SOFT", "At least 2 of S1, S2, S3 are true while in BASE_VALID, MARKUP_ACTIVE, or AVOID_SOFT.", "pivots, EMA30, EMA30 slope, OBV", "Blocks new entries; clears after 5 sessions without active soft condition."],
            ["AVOID_HARD", "H1 close below last markup swing low, or H2 two closes below originating/base top.", "significant low pivot, base_top_ref, close", "Forces open position to close with EXIT_AVOID_HARD."],
            ["MARKDOWN", "After AVOID_HARD if close is below EMA30.", "close and EMA30", "Risk-off state; returns to NEUTRAL after 5 closes above EMA30."],
        ],
    )

    add_heading(doc, "Core Calculations", 1)
    add_body(doc, "These are the calculations that directly affect lifecycle and exit behavior. Values are evaluated per symbol per trading day.")
    add_diagram(doc, diagrams[8], "Figure 9. Key formulas used by the state machine and harness.")
    add_formula(doc, "base_mfe = max(0, high / base_top_ref - 1)")
    add_formula(doc, "MARKUP_ACTIVE when base_state == BASE_RETIRED and base_mfe >= 0.20")
    add_formula(doc, "pivot significant when abs(pivot_price - opposite_pivot_price) >= 1.5 * ATR14")
    add_formula(doc, "position_mfe = max(previous_mfe, close / entry_close - 1)")
    add_formula(doc, "pnl_pct = (exit_close / entry_close - 1) * 100")
    add_formula(doc, "Variant B trail = max_close_since_entry - 2.75 * ATR14")

    add_heading(doc, "Entries and Exits", 1)
    add_body(doc, "Entries are shared between both variants. A position can open from a confirmed valid base, or during markup after a controlled pullback or flag breakout. Exits are partly shared and partly variant-specific. Both variants exit immediately on AVOID_HARD and both use the progress time-stop. The difference is the final open-position exit policy.")
    add_diagram(doc, diagrams[3], "Figure 4. Entry paths and exit checks.")

    add_heading(doc, "Entry Trigger Details", 2)
    add_rule_table(
        doc,
        ["Entry Path", "Trigger", "All Conditions"],
        [
            ["BASE_CONFIRMED_DIRECT", "Enter from BASE_VALID.", "lifecycle_state == BASE_VALID; candidate_intent_state == INTENT_FORMED; confirmation_state == CONFIRMED; base_valid is true; avoid_soft and avoid_hard are false; no existing position."],
            ["MARKUP_PULLBACK_EMA_BAND", "Enter during MARKUP_ACTIVE after pullback and recovery.", "lifecycle_state == MARKUP_ACTIVE; low touches between EMA10 and EMA30 or low <= EMA30; recovery occurs within 3 sessions; close >= EMA10; avoid tiers are clear; no existing position."],
            ["MARKUP_FLAG_BREAKOUT", "Enter during MARKUP_ACTIVE on compact-range breakout.", "lifecycle_state == MARKUP_ACTIVE; recent 5-15 session range width <= 2 x ATR14; current close breaks above that range high; avoid tiers are clear; no existing position."],
        ],
    )

    add_heading(doc, "Avoid Trigger Details", 2)
    add_rule_table(
        doc,
        ["Avoid Signal", "Calculation", "Meaning"],
        [
            ["S1_LOWER_HIGH", "last significant high price < prior significant high price", "Structure is making a lower high."],
            ["S2_TREND_LOSS", "close < EMA30 and EMA30 slope < 0", "Price is below the medium trend and the trend is falling."],
            ["S3_FLOW_DIVERGENCE", "last high >= prior high but OBV at last high < OBV at prior high", "Price matched or exceeded the old high, but participation weakened."],
            ["AVOID_SOFT", "count(S1,S2,S3 true) >= 2", "Warning tier. New entries are blocked; open positions are not forced out solely by soft tier."],
            ["H1_CLOSE_BELOW_MARKUP_SWING_LOW", "close < last_markup_swing_low", "Hard structural break below the last significant markup low."],
            ["H2_TWO_CLOSES_BELOW_BASE_TOP", "two-session streak where close < originating_base_top/base_top_ref", "Failed breakout / loss of base support area."],
            ["AVOID_HARD", "H1 or H2 true", "Hard risk-off tier. Any open position is closed immediately."],
        ],
    )

    add_heading(doc, "Shared Exit Details", 2)
    add_rule_table(
        doc,
        ["Exit", "Trigger", "Exit Reason"],
        [
            ["Hard avoid forced exit", "AVOID_HARD becomes true while a position is open.", "EXIT_AVOID_HARD"],
            ["Progress time-stop", "sessions_held >= 60 and then every 20 sessions; position MFE < 8%; lifecycle_state is not MARKUP_ACTIVE.", "EXITED_TIMESTOP_STAGNANT"],
            ["Time-stop waiver", "At check time, MFE >= 8% or lifecycle_state == MARKUP_ACTIVE.", "No exit; keep holding."],
        ],
    )

    add_heading(doc, "Variant A vs Variant B", 1)
    add_body(doc, "Variant A is structural. It asks whether the medium-term trend has broken. If price closes below EMA30 for two consecutive sessions, the position exits with EXIT_STRUCTURAL_EMA30_2C.")
    add_body(doc, "Variant B is trailing-stop based. It asks whether price has fallen too far from the best close since entry after adjusting for volatility. It exits when close is below max close since entry minus 2.75 times ATR14, with exit reason EXIT_CHANDELIER.")
    add_diagram(doc, diagrams[4], "Figure 5. The only intended difference between Variant A and Variant B.")

    add_rule_table(
        doc,
        ["Variant", "Exit Calculation", "Behavioral Intent", "Tradeoff"],
        [
            ["A", "ema30_loss_streak increments when close < EMA30; exit when streak >= 2.", "Stay while the medium trend structure is intact.", "More tolerant of volatility, but can give back more if price drops fast."],
            ["B", "trail = max_close_since_entry - 2.75 * ATR14; exit when close < trail.", "Protect progress after a move using volatility-adjusted distance.", "Can protect gains sooner, but can be shaken out if ATR does not absorb normal volatility."],
        ],
    )

    add_heading(doc, "Evidence Package", 1)
    add_body(doc, "When the full run completes, the harness exports daily files for SANAM, TIJARA, and MABANEE; universe summaries for both variants; a comparison against the v4.1-B baseline; an audit report; a gate report; and SHA256 sidecars for evidence integrity.")
    add_diagram(doc, diagrams[5], "Figure 6. Evidence and audit outputs produced after replay.")

    add_heading(doc, "Evidence Fields and What They Prove", 2)
    add_rule_table(
        doc,
        ["Evidence", "Purpose"],
        [
            ["v5A_sanam_daily.txt / v5B_sanam_daily.txt", "Shows daily lifecycle, avoid tier, confirmation, entry/exit, position state, and exit reason for SANAM."],
            ["v5A_tijara_daily.txt / v5B_tijara_daily.txt", "Shows TIJARA daily behavior and May shakeout evidence."],
            ["v5A_mabanee_daily.txt / v5B_mabanee_daily.txt", "Shows MABANEE avoid-tier changes and risk state behavior."],
            ["v5A_universe.txt / v5B_universe.txt", "Per-symbol and global counts: positions, exits, time-stops, avoid-hard exits, open sessions, best/worst/sum PnL."],
            ["v5_vs_baseline.txt", "Compares each candidate variant against the v4.1-B baseline."],
            ["v5_audit_report.txt", "Contains D4 evidence, position event chains, MABANEE avoid changes, and self-audit facts."],
            ["v5_gate_report.txt", "Reports G1-G7 gates: sealed SHA, universe assertions, freeze SHA, row parity, sidecars, unit checks, and constants."],
            ["*.sha256", "Hash sidecars proving exported file bytes."],
        ],
    )

    add_heading(doc, "Glossary", 1)
    add_bullets(
        doc,
        [
            "Sealed surface: the fixed input database used for reproducible replay.",
            "Ratified engine: an existing trusted service used as a black box by the harness.",
            "SessionContext: the daily packet of facts passed into the pure candidate state machine.",
            "Base top: the upper reference of a valid base, used for base MFE and hard breach checks.",
            "Significant pivot: a lag-confirmed high or low whose price distance from the opposite pivot is at least 1.5 x ATR14.",
            "EMA30: 30-session exponential moving average, used by Variant A as a trend boundary.",
            "EMA10: 10-session exponential moving average, used for markup pullback recovery entry.",
            "ATR14: 14-session average true range, used by Variant B to scale the trailing stop by volatility.",
            "OBV: on-balance volume, used to detect flow divergence at significant high pivots.",
            "MFE: maximum favorable excursion, the best progress a position has made since entry.",
            "AVOID_SOFT: warning tier that blocks new entries while conditions deteriorate.",
            "AVOID_HARD: hard risk-off tier that forces any open position to close.",
        ],
    )

    add_heading(doc, "Current Run Interpretation", 1)
    add_body(doc, "During the live run, seeing only a growing harness_v5A database means Variant A is still replaying. Variant B starts only after Variant A completes. The export folder appears only after both variants finish and the evidence package is written.")

    doc.save(DOCX_PATH)


def main() -> None:
    diagrams = [
        diagram_architecture(),
        diagram_daily_flow(),
        diagram_lifecycle(),
        diagram_entry_exit(),
        diagram_variants(),
        diagram_exports(),
        diagram_indicators(),
        diagram_trigger_matrix(),
        diagram_calculations(),
    ]
    build_doc(diagrams)
    print(DOCX_PATH)
    for path in diagrams:
        print(path)


if __name__ == "__main__":
    main()